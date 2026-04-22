"""
AI Courtroom v2.0 — Report Generation Service.

Generates real PDF and DOCX audit reports from analysis, courtroom,
and remediation data.  Uses:
  - reportlab for PDF (compatible, no system deps)
  - python-docx for DOCX

No WeasyPrint (requires GTK system deps).  reportlab is pure-Python and
production-ready.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Any

logger = logging.getLogger("courtroom.reports")


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF Generation (reportlab)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_pdf_report(
    session: dict,
    bias_metrics: list[dict],
    verdict: dict | None = None,
    remediation: dict | None = None,
) -> bytes:
    """Generate a comprehensive PDF audit report.  Returns raw bytes."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=22, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceBefore=14, spaceAfter=8)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, spaceBefore=10, spaceAfter=6)
    body = styles["Normal"]
    small = ParagraphStyle("Small", parent=body, fontSize=8, textColor=colors.grey)

    elements: list = []

    # ── Title page ──
    elements.append(Spacer(1, 40*mm))
    elements.append(Paragraph("AI COURTROOM", title_style))
    elements.append(Paragraph("Bias Audit Report", styles["Heading2"]))
    elements.append(Spacer(1, 10*mm))
    elements.append(HRFlowable(width="60%", thickness=2, color=colors.HexColor("#f59e0b")))
    elements.append(Spacer(1, 10*mm))

    # Session info
    info_data = [
        ["Session ID", session.get("session_id", "N/A")],
        ["Dataset", session.get("dataset_filename", "N/A")],
        ["Model", session.get("model_filename", "N/A")],
        ["Rows", str(session.get("row_count", "N/A"))],
        ["Features", str(session.get("feature_count", "N/A"))],
        ["Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")],
    ]
    t = Table(info_data, colWidths=[120, 300])
    t.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#6b7280")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)
    elements.append(PageBreak())

    # ── Bias Metrics ──
    elements.append(Paragraph("1. Fairness Metrics", h2))
    elements.append(Paragraph(
        "The following metrics were computed using Fairlearn against the uploaded dataset and model.",
        body,
    ))
    elements.append(Spacer(1, 4*mm))

    if bias_metrics:
        metric_rows = [["Metric", "Value", "Threshold", "Status", "Severity"]]
        for m in bias_metrics:
            sev = m.get("severity", "unknown")
            if sev == "pass_":
                sev = "pass"
            metric_rows.append([
                m.get("metric_name", ""),
                f'{m.get("metric_value", 0):.4f}',
                str(m.get("threshold", "")),
                "PASS" if m.get("passed") else "FAIL",
                sev.upper(),
            ])

        mt = Table(metric_rows, colWidths=[150, 70, 70, 50, 70])
        mt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#334155")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ]))

        # Color-code severity
        for i, m in enumerate(bias_metrics, start=1):
            sev = m.get("severity", "")
            if sev == "critical":
                mt.setStyle(TableStyle([("TEXTCOLOR", (4, i), (4, i), colors.red)]))
            elif sev == "warning":
                mt.setStyle(TableStyle([("TEXTCOLOR", (4, i), (4, i), colors.HexColor("#f59e0b"))]))
            else:
                mt.setStyle(TableStyle([("TEXTCOLOR", (4, i), (4, i), colors.green)]))

            if not m.get("passed"):
                mt.setStyle(TableStyle([("TEXTCOLOR", (3, i), (3, i), colors.red)]))

        elements.append(mt)
    else:
        elements.append(Paragraph("<i>No bias metrics available.</i>", body))

    # ── Courtroom Verdict ──
    if verdict:
        elements.append(Spacer(1, 6*mm))
        elements.append(Paragraph("2. Courtroom Verdict", h2))

        v = verdict.get("judge_verdict", "unknown")
        score = verdict.get("bias_risk_score", 0)
        verdict_color = colors.red if v == "guilty" else colors.green

        elements.append(Paragraph(f"<b>Verdict:</b> {v.upper()}", body))
        elements.append(Paragraph(f"<b>Bias Risk Score:</b> {score}/100", body))
        elements.append(Spacer(1, 3*mm))

        elements.append(Paragraph("Prosecution Argument", h3))
        elements.append(Paragraph(verdict.get("prosecution_argument", "N/A"), body))

        elements.append(Paragraph("Defense Argument", h3))
        elements.append(Paragraph(verdict.get("defense_argument", "N/A"), body))

        elements.append(Paragraph("Judge Reasoning", h3))
        elements.append(Paragraph(verdict.get("judge_reasoning", "N/A"), body))

        elements.append(Paragraph("Recommended Sentence", h3))
        elements.append(Paragraph(verdict.get("recommended_sentence", "N/A"), body))

    # ── Remediation ──
    if remediation and remediation.get("strategy"):
        elements.append(Spacer(1, 6*mm))
        section_num = "3" if verdict else "2"
        elements.append(Paragraph(f"{section_num}. Remediation Results", h2))

        rem_data = [
            ["Strategy", remediation.get("strategy", "N/A")],
            ["Original Accuracy", f'{remediation.get("original_accuracy", 0):.4f}'],
            ["Mitigated Accuracy", f'{remediation.get("mitigated_accuracy", 0):.4f}'],
            ["Original DIR", f'{remediation.get("original_dir", 0):.4f}'],
            ["Mitigated DIR", f'{remediation.get("mitigated_dir", 0):.4f}'],
        ]
        rt = Table(rem_data, colWidths=[140, 200])
        rt.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(rt)

    # ── Footer ──
    elements.append(Spacer(1, 20*mm))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
    elements.append(Spacer(1, 2*mm))
    elements.append(Paragraph(
        "Generated by AI Courtroom v2.0 — Powered by Fairlearn, SHAP, and Claude",
        small,
    ))

    doc.build(elements)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX Generation (python-docx)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_docx_report(
    session: dict,
    bias_metrics: list[dict],
    verdict: dict | None = None,
    remediation: dict | None = None,
) -> bytes:
    """Generate a comprehensive DOCX audit report.  Returns raw bytes."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading("AI COURTROOM — Bias Audit Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    # Session info table
    doc.add_heading("Session Information", level=2)
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"
    info_pairs = [
        ("Session ID", session.get("session_id", "N/A")),
        ("Dataset", session.get("dataset_filename", "N/A")),
        ("Model", session.get("model_filename", "N/A")),
        ("Rows", str(session.get("row_count", "N/A"))),
        ("Features", str(session.get("feature_count", "N/A"))),
        ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
    ]
    for i, (label, value) in enumerate(info_pairs):
        t.cell(i, 0).text = label
        t.cell(i, 1).text = value
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")

    # Bias Metrics
    doc.add_heading("1. Fairness Metrics", level=2)
    if bias_metrics:
        mt = doc.add_table(rows=len(bias_metrics) + 1, cols=5)
        mt.style = "Table Grid"
        headers = ["Metric", "Value", "Threshold", "Status", "Severity"]
        for j, h in enumerate(headers):
            mt.cell(0, j).text = h
            for para in mt.cell(0, j).paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for i, m in enumerate(bias_metrics, start=1):
            sev = m.get("severity", "unknown")
            if sev == "pass_":
                sev = "pass"
            mt.cell(i, 0).text = m.get("metric_name", "")
            mt.cell(i, 1).text = f'{m.get("metric_value", 0):.4f}'
            mt.cell(i, 2).text = str(m.get("threshold", ""))
            mt.cell(i, 3).text = "PASS" if m.get("passed") else "FAIL"
            mt.cell(i, 4).text = sev.upper()
    else:
        doc.add_paragraph("No bias metrics available.")

    # Courtroom Verdict
    if verdict:
        doc.add_heading("2. Courtroom Verdict", level=2)
        v = verdict.get("judge_verdict", "unknown")
        p = doc.add_paragraph()
        run = p.add_run(f"Verdict: {v.upper()}")
        run.bold = True
        run.font.size = Pt(14)
        if v == "guilty":
            run.font.color.rgb = RGBColor(220, 38, 38)
        else:
            run.font.color.rgb = RGBColor(16, 185, 129)

        doc.add_paragraph(f"Bias Risk Score: {verdict.get('bias_risk_score', 0)}/100")

        doc.add_heading("Prosecution Argument", level=3)
        doc.add_paragraph(verdict.get("prosecution_argument", "N/A"))

        doc.add_heading("Defense Argument", level=3)
        doc.add_paragraph(verdict.get("defense_argument", "N/A"))

        doc.add_heading("Judge Reasoning", level=3)
        doc.add_paragraph(verdict.get("judge_reasoning", "N/A"))

        doc.add_heading("Recommended Sentence", level=3)
        doc.add_paragraph(verdict.get("recommended_sentence", "N/A"))

    # Remediation
    if remediation and remediation.get("strategy"):
        section_num = "3" if verdict else "2"
        doc.add_heading(f"{section_num}. Remediation Results", level=2)

        rt = doc.add_table(rows=5, cols=2)
        rt.style = "Table Grid"
        rem_pairs = [
            ("Strategy", remediation.get("strategy", "N/A")),
            ("Original Accuracy", f'{remediation.get("original_accuracy", 0):.4f}'),
            ("Mitigated Accuracy", f'{remediation.get("mitigated_accuracy", 0):.4f}'),
            ("Original DIR", f'{remediation.get("original_dir", 0):.4f}'),
            ("Mitigated DIR", f'{remediation.get("mitigated_dir", 0):.4f}'),
        ]
        for i, (label, value) in enumerate(rem_pairs):
            rt.cell(i, 0).text = label
            rt.cell(i, 1).text = value

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph("Generated by AI Courtroom v2.0 — Powered by Fairlearn, SHAP, and Claude")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(107, 114, 128)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
